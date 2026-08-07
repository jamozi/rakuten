from __future__ import annotations

import html
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "doc"
DOCX_PATH = OUTPUT_DIR / "raos_human_work_support_pack_ja-JP.docx"
HTML_PATH = ROOT / "tmp" / "docs" / "raos_human_work_support_pack_ja-JP.html"

NAVY = "17324D"
BLUE = "2B6CB0"
TEAL = "0F766E"
AMBER = "A16207"
RED = "B42318"
INK = "1F2937"
MUTED = "5B6673"
PALE_BLUE = "EAF2FB"
PALE_TEAL = "E8F5F2"
PALE_AMBER = "FFF6DB"
PALE_GRAY = "F3F5F7"
WHITE = "FFFFFF"
FONT = "Noto Sans CJK JP"


DECISIONS = [
    {
        "id": "OD-001",
        "phase": "最初に必要",
        "topic": "初期カテゴリ",
        "owner": "Product Owner",
        "human": "低リスクで比較軸を構造化できる最初の商品カテゴリを1つ選ぶ。",
        "default": "カテゴリ固有実装を止め、合成Fixtureだけを使う。",
        "support": "候補の採点、規約・需要・更新負荷・収益性の比較、Decision Record案を作る。",
        "done": "カテゴリ名、対象外範囲、採用理由、承認者、承認日が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-002",
        "phase": "最初に必要",
        "topic": "サイト名・ドメイン",
        "owner": "Product Owner",
        "human": "ブランド名、取得するドメイン、運営者表記を決め、ドメインを契約する。",
        "default": "仮ブランドとexample.invalidを使い、外部公開しない。",
        "support": "名称案、重複・商標・ドメイン候補調査、運営者表記とページ構成案を作る。",
        "done": "ブランド、ドメイン、所有者、登録媒体、運営者表記が承認されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-003",
        "phase": "Finance実装前",
        "topic": "楽天成果Reportサンプル",
        "owner": "Business Owner",
        "human": "楽天管理画面から実Reportを取得し、匿名化したサンプルと列・状態・粒度を確認する。",
        "default": "Synthetic fixtureのみを使い、実成果帰属は未検証と表示する。",
        "support": "匿名化手順、Secret/PII検査、Parser、列対応表、Recorded Contract Testを作る。",
        "done": "匿名化済みサンプルの参照先、Report期間、列定義、承認者が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-004",
        "phase": "GATE-2前",
        "topic": "順位・Keyword Provider",
        "owner": "Product Owner",
        "human": "規約適合するProviderを契約するか、Search Consoleと手動CSVで運用するか決める。",
        "default": "Search Consoleと手動CSVだけを使う。",
        "support": "候補比較、料金・規約・API・データ保持の評価、Adapter案を作る。",
        "done": "利用方式、許可データ、契約主体、費用上限が記録されている。",
        "blocking": "No",
    },
    {
        "id": "OD-005",
        "phase": "最初に必要",
        "topic": "Reviewerと人件費",
        "owner": "Business Owner",
        "human": "主Reviewer、代理Reviewer、役割分離、標準時間単価を決める。",
        "default": "公開を止め、人件費をUNKNOWNとして扱う。",
        "support": "役割表、レビュー手順、工数記録票、利益計算への反映を用意する。",
        "done": "2つの独立した承認主体、代理者、単価、稼働可能時間が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-006",
        "phase": "カテゴリ実装前",
        "topic": "商品同一性ルール",
        "owner": "Domain Editor",
        "human": "型番、容量、色、セット、JAN等を統合するか分離するかカテゴリ別に定義する。",
        "default": "自動統合せず、Human Reviewへ送る。",
        "support": "候補商品の差分表、曖昧ケース、Golden Product Fixture、Rule案を作る。",
        "done": "統合・分離ルール、例外、代表商品、Domain Reviewer承認がある。",
        "blocking": "Yes",
    },
    {
        "id": "OD-007",
        "phase": "カテゴリ実装前",
        "topic": "鮮度SLA",
        "owner": "Managing Editor",
        "human": "価格、在庫、仕様、画像、リンクの最大許容Ageを決める。",
        "default": "保守的な暫定値を使い、Stale時は非表示にする。",
        "support": "更新頻度、変動リスク、API制限、作業量からSLA案を比較する。",
        "done": "項目別日数、Stale時挙動、例外、承認者が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-008",
        "phase": "GATE-0前",
        "topic": "法務レビュー境界",
        "owner": "Business Owner / Legal",
        "human": "法務確認が必要な表現・カテゴリ・例外、相談経路、記録方法を決める。",
        "default": "AIや開発者が法的判断を代替せず、公開を止める。",
        "support": "論点一覧、広告・著作権・Privacyの確認票、相談記録テンプレートを作る。",
        "done": "確認対象、相談先、SLA、例外承認者、記録保管先が承認されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-009",
        "phase": "最初に必要",
        "topic": "予算・許容赤字",
        "owner": "Business Owner",
        "human": "AWS、LLM、外部Providerの月次上限、許容赤字期間、自動停止閾値を決める。",
        "default": "低い開発上限だけを使い、Productionを無効にする。",
        "support": "固定費・変動費・記事別費用のシナリオ、停止閾値、Alert案を作る。",
        "done": "月次上限、日次上限、1 Job上限、赤字期間、停止・再開承認者が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-010",
        "phase": "Admin公開前",
        "topic": "OIDC Provider",
        "owner": "Security Owner",
        "human": "Cognitoまたは承認済みOIDC Providerを選び、契約・管理主体を決める。",
        "default": "Local fake authは開発専用とし、外部公開しない。",
        "support": "比較表、IAM/RBAC/MFA設計、設定とテストを実装する。",
        "done": "Provider、管理者、MFA、復旧方法、費用、環境分離が承認されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-011",
        "phase": "Production前",
        "topic": "通知・Escalation",
        "owner": "Operations Owner",
        "human": "Critical/Highの通知先、一次対応者、代理者、連絡手段を決める。",
        "default": "Local logだけを使い、Productionへ進めない。",
        "support": "Alert routing、Runbook、当番表、Tabletopシナリオを作る。",
        "done": "一次・代理・Incident Commander候補、連絡先参照、応答目標がある。",
        "blocking": "Yes",
    },
    {
        "id": "OD-012",
        "phase": "Public analytics前",
        "topic": "Privacy・同意",
        "owner": "Privacy Owner",
        "human": "Cookie、計測、同意Banner、Privacy文言、拒否時挙動を承認する。",
        "default": "非必須Trackingを無効にし、First-party最小Eventだけを使う。",
        "support": "データフロー、収集最小化案、文言下書き、Consent実装を作る。",
        "done": "収集項目、目的、同意方式、拒否時挙動、第三者提供、承認者が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-013",
        "phase": "Terraform production前",
        "topic": "Region・Data Residency",
        "owner": "Security / Business Owner",
        "human": "本番Region、Backup Region、越境移転の扱いを承認する。",
        "default": "ap-northeast-1をReferenceとし、Production applyを禁止する。",
        "support": "構成・費用・RTO/RPO・データ移転の比較とIaC案を作る。",
        "done": "Primary/Backup Region、対象データ、越境方針、承認者が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-014",
        "phase": "Deletion jobs前",
        "topic": "保持期間",
        "owner": "Privacy / Finance / Legal",
        "human": "Analytics個票、Security Log、AI Artifact、成果データの保持期間を承認する。",
        "default": "自動削除を無効にし、収集を最小化する。",
        "support": "データ棚卸し、保持期間表、Legal hold、削除・検証Jobを作る。",
        "done": "データ種別別の期間、根拠、削除方法、例外、承認者が記録されている。",
        "blocking": "Yes",
    },
    {
        "id": "OD-015",
        "phase": "Live adapter前",
        "topic": "本番Provider資格情報",
        "owner": "Operations Owner",
        "human": "楽天、OpenAI、Google、AWSの専用Account、権限、Secretを安全に設定する。",
        "default": "Recorded fixtureだけを使う。",
        "support": "最小権限IAM、Secret参照、Rotation手順、Live smokeを実装する。",
        "done": "Account owner、権限、Secret参照名、Rotation、失効手順が記録されている。",
        "blocking": "Yes",
    },
]


PROCEDURES = [
    ("事業主体", "個人/法人の運営主体、責任者、会計方針を確定", "Business Owner", "未着手"),
    ("楽天Affiliate", "加入・規約同意・本人情報・入金先・媒体URL登録", "Business Owner", "未着手"),
    ("楽天Web Service", "Application登録、利用条件確認、専用Credential発行", "Operations Owner", "未着手"),
    ("ドメイン", "取得、更新責任者、DNS/HTTPS、登録媒体との一致確認", "Product Owner", "未着手"),
    ("AWS", "本番専用Account、請求、Budget、MFA、管理者復旧", "Security/Operations", "未着手"),
    ("OpenAI", "専用Project、請求上限、Service Account、データ条件確認", "Operations Owner", "未着手"),
    ("Google", "Search Console/Analytics Property、権限、同意方針", "Privacy/Operations", "未着手"),
    ("法務・Privacy", "相談先、表示文、Copyright、Cookie、保持期間の承認", "Legal/Privacy", "未着手"),
    ("税務・会計", "申告主体、売上・経費・外注費・証憑の扱いを確認", "Business Owner", "未着手"),
]


OPERATIONS = [
    ("毎日", "Critical Alert、Queue/DLQ、リンク、鮮度、Cost velocity", "Operations Owner / 代理"),
    ("毎週", "Provider失敗、Stale queue、Security finding、Policy期限", "Operations + Managing Editor"),
    ("毎月", "権限、確定成果、原価、Backup、Retention、SLO", "Business + Security + Operations"),
    ("四半期", "Restore drill、Incident tabletop、Runbook、Threat review", "全Owner"),
    ("記事ごと", "Fact、商品同定、推薦、Policy、Preview差分、最終公開", "Editor + Reviewer + Approver"),
    ("Releaseごと", "Contract、Migration、Security、Evidence、Rollback、承認", "Engineering + Security + Operations + PO"),
]


SOURCES = [
    "docs/canonical/01_integration/RAOS_07_open_decisions_v1.0.yaml",
    "docs/canonical/04_security/RAOS_10_role_permission_matrix_v1.0.yaml",
    "docs/canonical/05_test/RAOS_11_release_evidence_template_v1.0.yaml",
    "docs/canonical/06_ops/RAOS_12_operations_reliability_design_v1.0.md",
    "docs/canonical/07_backlog/RAOS_13_story_backlog_v1.0.csv",
    "docs/upstream/key_documents/RAOS_01_requirements_purpose_success_v0.1.md",
    "docs/upstream/key_documents/RAOS_05_ai_agent_prompt_routing_evaluation_design_v0.1.md",
    "docs/upstream/key_documents/RAOS_06_content_editorial_evidence_design_v0.1.md",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)
    keep_next.set(qn("w:val"), "1" if value else "0")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("RAOS Human Work Support Pack  |  ")
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for style_name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 12, MUTED, 0, 8),
        ("Heading 1", 18, NAVY, 12, 7),
        ("Heading 2", 13, BLUE, 9, 5),
        ("Heading 3", 10.5, TEAL, 6, 3),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "RAOS  |  人手作業支援パック  |  非Canonical支援資料"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_text(doc: Document, text: str, bold=False, color: str | None = None, size: float | None = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc: Document, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = p.add_run(text)
    run.font.name = FONT
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = FONT
    return p


def add_callout(doc: Document, title: str, body: str, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=190, bottom=150, end=190)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = FONT
    r.font.color.rgb = RGBColor.from_string(accent)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    return table


def style_table(table, header_fill=NAVY, font_size=8.2):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if table.rows:
        set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)


def add_field_table(doc: Document, fields: list[tuple[str, str]], widths=(4.2, 13.0)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(widths[0])
    table.columns[1].width = Cm(widths[1])
    for idx, (label, value) in enumerate(fields):
        cells = table.add_row().cells
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], PALE_BLUE)
        for cell in cells:
            set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = Pt(8.5)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    return table


def add_decision_sheet(doc: Document, item: dict, page_break=False):
    if page_break:
        doc.add_page_break()
    heading = doc.add_heading(f"{item['id']}  {item['topic']}", level=2)
    set_keep_with_next(heading)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    values = [
        ("必要時点", item["phase"], "Owner", item["owner"]),
        ("Production blocker", item["blocking"], "現在の状態", "未決定"),
    ]
    for row, data in zip(table.rows, values):
        for idx, text in enumerate(data):
            row.cells[idx].text = text
            set_cell_margins(row.cells[idx])
            if idx in (0, 2):
                set_cell_shading(row.cells[idx], NAVY)
                for run in row.cells[idx].paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            else:
                set_cell_shading(row.cells[idx], PALE_GRAY)
            for p in row.cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = Pt(8.2)
    add_field_table(
        doc,
        [
            ("人が決める・提供する", item["human"]),
            ("未決定時の安全動作", item["default"]),
            ("Codexが支援できること", item["support"]),
            ("完了条件", item["done"]),
            ("決定・提供内容", "\n\n"),
            ("理由・比較した選択肢", "\n\n"),
            ("制約・例外・有効期限", "\n"),
            ("承認者・承認日", "承認者: ____________________    日付: ____ / ____ / ____"),
            ("Evidence参照", "ファイル/URL/チケット参照のみ。Secret値は記載しない: ____________________"),
        ],
    )


def add_form_title(doc: Document, title: str, subtitle: str):
    doc.add_page_break()
    doc.add_heading(title, level=1)
    add_text(doc, subtitle, color=MUTED)


def build_docx() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    r = p.add_run("RAOS")
    r.font.name = FONT
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    title = doc.add_paragraph(style="Title")
    title.add_run("人手作業支援パック")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("決める・提供する・承認する作業を、Codexと分担して前へ進める")
    add_callout(
        doc,
        "この資料の位置づけ",
        "現行Canonicalに基づく記入用の非Canonical支援資料です。記入しただけではOpen Decisionは解決されません。人間の承認、Canonicalとの照合、所定のDecision RecordまたはEvidence登録が完了して初めて実装・本番判断へ使えます。",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    add_text(doc, "Version 1.0  |  2026-08-04  |  ja-JP", color=MUTED, size=9)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("重要")
    r.bold = True
    r.font.name = FONT
    r.font.color.rgb = RGBColor.from_string(RED)
    p.add_run("  Secret、Password、Token、個人情報の原文は、この資料やチャットへ記載しないでください。")

    doc.add_page_break()
    doc.add_heading("1. 使い方", level=1)
    add_number(doc, "人が事業・法務・契約・公開に関する選択を行います。")
    add_number(doc, "Codexが候補比較、リスク整理、設定案、チェックリスト、テストを作ります。")
    add_number(doc, "人がDecision Recordまたは承認票を確認し、承認します。")
    add_number(doc, "Codexが承認済み範囲だけを実装し、Local/CI/Runtime/Production証跡を分離して報告します。")
    add_callout(
        doc,
        "境界",
        "Codexは、法的判断、規約同意、本人確認、契約締結、支払い承認、Secret値の保管、記事の最終公開、成果確定、Kill Switch解除を人間の代わりに行いません。",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    doc.add_heading("最初の30分で回答する5項目", level=2)
    kickoff = [
        ("1. 初期カテゴリ", "決定済み: __________  /  候補3つ: ______________________________"),
        ("2. ブランド・ドメイン", "名称案: __________________  ドメイン候補: __________________"),
        ("3. 予算・許容赤字", "月額上限: ______円  許容赤字期間: ____か月  自動停止: ______円"),
        ("4. Reviewer", "主担当: __________  代理: __________  時間単価: ______円"),
        ("5. 外部手続き", "楽天: 未/済  Domain: 未/済  AWS: 未/済  Legal相談先: 未/済"),
    ]
    add_field_table(doc, kickoff)
    add_text(doc, "チャット回答テンプレート", bold=True, color=NAVY)
    add_callout(
        doc,
        "コピーして回答できます",
        "初期カテゴリ候補: / ブランド名候補: / ドメイン候補: / 月額予算上限: / 許容赤字期間: / 主Reviewer: / 代理Reviewer: / 標準時間単価: / 楽天Affiliate登録状況: / 法務相談先の有無:",
        fill=PALE_GRAY,
        accent=NAVY,
    )

    doc.add_page_break()
    doc.add_heading("2. Open Decision全体表", level=1)
    add_text(doc, "15件のうち14件がProductionまたはGate blockerです。OD-004だけは手動Importで先へ進められます。")
    table = doc.add_table(rows=1, cols=6)
    headers = ["ID", "必要時点", "判断・Evidence", "Owner", "Block", "状態"]
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
    for item in DECISIONS:
        cells = table.add_row().cells
        for cell, value in zip(
            cells,
            [item["id"], item["phase"], item["topic"], item["owner"], item["blocking"], "未決定"],
        ):
            cell.text = value
    widths = [1.5, 2.6, 4.0, 4.1, 1.4, 2.0]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Cm(width)
    style_table(table, font_size=7.5)

    doc.add_heading("進行順序", level=2)
    add_bullet(doc, "A. 今決める: OD-001、002、005、008、009")
    add_bullet(doc, "B. カテゴリ設計時: OD-006、007")
    add_bullet(doc, "C. Finance/Live Adapter時: OD-003、015")
    add_bullet(doc, "D. Staging/Production前: OD-010、011、012、013、014")
    add_bullet(doc, "E. GATE-2高度化時: OD-004。未決でもSearch Consoleと手動CSVで継続可能")

    doc.add_page_break()
    doc.add_heading("3. Decision記入シート", level=1)
    add_text(doc, "各シートを人が記入・承認し、CodexがCanonicalと照合して正式なDecision/Evidence候補へ変換します。")
    for index, item in enumerate(DECISIONS):
        add_decision_sheet(doc, item, page_break=index > 0)

    add_form_title(doc, "4. 外部手続きチェックリスト", "契約・本人確認・支払い・媒体登録はAccount Owner本人が行います。Evidence参照には確認番号やファイル参照だけを記録し、Secret値を記載しません。")
    table = doc.add_table(rows=1, cols=6)
    headers = ["完了", "対象", "人が行うこと", "Owner", "状態", "Evidence参照"]
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
    for name, task, owner, status in PROCEDURES:
        cells = table.add_row().cells
        for cell, value in zip(cells, ["☐", name, task, owner, status, ""]):
            cell.text = value
    style_table(table, font_size=7.6)
    add_callout(
        doc,
        "Credentialの扱い",
        "Codexへ渡すのはSecret値ではなく、RAOS_RAKUTEN_ACCESS_KEYなどの参照名、権限範囲、有効期限、Rotation状態です。漏えいが疑われる場合、ローカル削除とProvider側失効は別作業として両方確認します。",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    add_form_title(doc, "5. 記事レビュー・公開承認票", "CodexはDraft、Evidence照合、Finding、Preview差分を準備します。ReviewerとApproverが内容を確認し、公開可否を決めます。")
    add_field_table(
        doc,
        [
            ("Article ID / Version", ""),
            ("カテゴリ / Primary Decision", ""),
            ("Reviewer / Approver", "Reviewer: __________________  Approver: __________________"),
            ("Evidence", "☐ 主要Claim 100%  ☐ 全検証可能Claim 95%以上  ☐ Source原本を確認"),
            ("商品同定", "☐ 型番  ☐ 容量  ☐ 色  ☐ セット  ☐ JAN  ☐ 曖昧候補なし/処理済み"),
            ("品質", "☐ 85点以上  ☐ 各軸Floor合格  ☐ 架空体験なし  ☐ Tradeoff/対象外を明記"),
            ("Policy", "☐ Blocking Finding 0  ☐ 広告表示  ☐ 楽天遷移明示  ☐ 画像/引用許可"),
            ("鮮度・リンク", "☐ 価格  ☐ 在庫  ☐ 仕様  ☐ 画像  ☐ Affiliate linkが許容Age内"),
            ("Preview差分", "☐ 承認対象Version/Hashと一致  ☐ 公開Renderer差分確認"),
            ("判断", "☐ APPROVE  ☐ REJECT  ☐ PAUSE  理由: ______________________________"),
            ("承認", "Reviewer署名/ID: ______________ 日付: ______  Approver署名/ID: ______________ 日付: ______"),
        ],
    )
    add_callout(doc, "役割分離", "編集者本人だけで最終承認・公開を完結させません。Critical AI TaskのReleaseは二人承認を要求します。", fill=PALE_TEAL, accent=TEAL)

    add_form_title(doc, "6. 月次Finance Review票", "発生報酬ではなく、楽天の確定成果と実費を用いて月次貢献利益を承認します。")
    add_field_table(
        doc,
        [
            ("対象月 / Report期間", ""),
            ("Report Evidence", "取得元・ファイルHash・取込ID: ____________________________________"),
            ("照合", "☐ 発生/確定を分離  ☐ Cancel/否認/調整を確認  ☐ 未帰属を分離  ☐ 重複なし"),
            ("確定成果報酬", "________________ 円"),
            ("人件費", "________________ 円  根拠: ________________________________"),
            ("AI/API/Cloud費", "________________ 円"),
            ("その他直接費", "________________ 円"),
            ("確定貢献利益", "________________ 円"),
            ("未解決差異", "金額: __________円  理由/調査Owner: ____________________________"),
            ("判断", "☐ ACCEPT  ☐ HOLD  ☐ REIMPORT  ☐ INVESTIGATE"),
            ("承認", "Finance Reviewer: __________________  日付: ____ / ____ / ____"),
        ],
    )

    add_form_title(doc, "7. Release・Production承認票", "Local PASS、CI、Staging Runtime、Production readinessを混同せず、各OwnerがEvidenceを確認します。")
    add_field_table(
        doc,
        [
            ("Release ID / Commit", ""),
            ("対象環境 / Scope", ""),
            ("Test Evidence", "☐ Contract  ☐ Unit  ☐ Integration  ☐ E2E  ☐ Security  ☐ Manual"),
            ("Operations", "☐ Rollback tested  ☐ Backup restore current  ☐ Alerts tested  ☐ Runbooks current"),
            ("Open Decisions", "☐ Blocking Open Decision 0  未解決参照: __________________________"),
            ("Engineering", "☐ APPROVE  ☐ REJECT  Approver/Date: __________________________"),
            ("Security", "☐ APPROVE  ☐ REJECT  Approver/Date: __________________________"),
            ("Operations", "☐ APPROVE  ☐ REJECT  Approver/Date: __________________________"),
            ("Product Owner", "☐ APPROVE  ☐ REJECT  Approver/Date: __________________________"),
            ("最終判断", "☐ NOT_READY  ☐ CANARY  ☐ PRODUCTION  ☐ ROLLBACK"),
        ],
    )

    add_form_title(doc, "8. Incident・Kill Switch承認票", "誤公開、Credential侵害、Affiliate不正、重大品質事故ではTraffic維持より安全停止を優先します。")
    add_field_table(
        doc,
        [
            ("Incident ID / Severity", ""),
            ("検知日時 / 宣言者", ""),
            ("影響", "公開 / Affiliate / Data / Credential / Finance / その他: __________________"),
            ("初動", "☐ 公開Kill  ☐ Affiliate Kill  ☐ Credential失効  ☐ Rollback  ☐ Provider連絡"),
            ("Evidence", "Timeline/Log/Artifact参照: __________________________________________"),
            ("復旧条件", "_______________________________________________________________"),
            ("再開前確認", "☐ Root cause  ☐ Corrective action  ☐ Regression test  ☐ Owner sign-off"),
            ("解除承認", "Product Owner/Operator: __________________  日付: ____ / ____ / ____"),
        ],
    )

    add_form_title(doc, "9. 継続運用カレンダー", "Codexが収集・集計・差分作成を自動化し、人間が意味判断と承認を行います。")
    table = doc.add_table(rows=1, cols=4)
    for cell, label in zip(table.rows[0].cells, ["周期", "人が確認すること", "主担当", "実施/次回"]):
        cell.text = label
    for cycle, task, owner in OPERATIONS:
        cells = table.add_row().cells
        for cell, value in zip(cells, [cycle, task, owner, ""]):
            cell.text = value
    style_table(table, font_size=8.0)

    doc.add_heading("最小の人員構成", level=2)
    add_bullet(doc, "Product/Business Owner: 事業目標、カテゴリ、予算、契約、Gate判断")
    add_bullet(doc, "Managing Editor/Reviewer: 商品同定、Evidence、記事品質、推薦、最終編集")
    add_bullet(doc, "Operations/Security Ownerと代理: Account、Credential、Alert、Incident、Recovery")
    add_bullet(doc, "必要時のLegal/Privacy/Tax専門家: 規約、表示、Copyright、Privacy、税務")
    add_callout(doc, "最低人数", "公開フローでは役割分離のため、少なくとも2つの独立した人間主体を確保します。本番運用では一次対応者と代理者も決めます。", fill=PALE_TEAL, accent=TEAL)

    doc.add_heading("Codexへの依頼テンプレート", level=2)
    add_callout(
        doc,
        "意思決定支援を依頼する",
        "OD-___について支援してください。候補/前提は____です。比較軸は____、月額上限は____円、避けたいリスクは____です。Canonicalと公式情報を確認し、推奨案、代替案、未解決点、Decision Record案を作ってください。Secret値は扱わないでください。",
        fill=PALE_GRAY,
        accent=NAVY,
    )

    doc.add_page_break()
    doc.add_heading("10. 参照元と注意事項", level=1)
    add_text(doc, "この支援パックは次のRepository内Artifactを要約しています。Canonical/Upstreamは編集せず、差異がある場合はCanonicalの優先順位に従います。")
    for source in SOURCES:
        add_bullet(doc, source)
    add_callout(
        doc,
        "完了の定義",
        "チェック欄が埋まったこと自体はProduction readinessではありません。承認済みDecision、正式Evidence、実Runtime結果、Security/Operations/Product OwnerのSign-offがそろって初めてGate判定できます。",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    add_text(doc, "End of document", color=MUTED, size=8)
    doc.save(DOCX_PATH)


def html_escape(value: str) -> str:
    return html.escape(value).replace("\n", "<br>")


def decision_html(item: dict) -> str:
    rows = [
        ("人が決める・提供する", item["human"]),
        ("未決定時の安全動作", item["default"]),
        ("Codexが支援できること", item["support"]),
        ("完了条件", item["done"]),
        ("決定・提供内容", "<div class='blank tall'></div>"),
        ("理由・比較した選択肢", "<div class='blank tall'></div>"),
        ("制約・例外・有効期限", "<div class='blank'></div>"),
        ("承認者・承認日", "承認者: ____________________　日付: ____ / ____ / ____"),
        ("Evidence参照", "ファイル/URL/チケット参照のみ。Secret値は記載しない: ____________________"),
    ]
    body = "".join(f"<tr><th>{html_escape(k)}</th><td>{v if v.startswith('<div') else html_escape(v)}</td></tr>" for k, v in rows)
    return f"""
    <section class="page decision">
      <h2>{item['id']}　{html_escape(item['topic'])}</h2>
      <table class="meta"><tr><th>必要時点</th><td>{html_escape(item['phase'])}</td><th>Owner</th><td>{html_escape(item['owner'])}</td></tr>
      <tr><th>Production blocker</th><td>{item['blocking']}</td><th>現在の状態</th><td>未決定</td></tr></table>
      <table class="fields">{body}</table>
    </section>"""


def build_html() -> None:
    summary_rows = "".join(
        f"<tr><td>{d['id']}</td><td>{html_escape(d['phase'])}</td><td>{html_escape(d['topic'])}</td><td>{html_escape(d['owner'])}</td><td>{d['blocking']}</td><td>未決定</td></tr>"
        for d in DECISIONS
    )
    decision_pages = "".join(decision_html(item) for item in DECISIONS)
    procedure_rows = "".join(
        f"<tr><td class='check'>☐</td><td>{html_escape(name)}</td><td>{html_escape(task)}</td><td>{html_escape(owner)}</td><td>{status}</td><td></td></tr>"
        for name, task, owner, status in PROCEDURES
    )
    operation_rows = "".join(
        f"<tr><td>{html_escape(cycle)}</td><td>{html_escape(task)}</td><td>{html_escape(owner)}</td><td></td></tr>"
        for cycle, task, owner in OPERATIONS
    )
    source_items = "".join(f"<li>{html_escape(s)}</li>" for s in SOURCES)
    css = f"""
      @page {{ size: A4; margin: 14mm 14mm 15mm 14mm; @bottom-right {{ content: counter(page); }} }}
      * {{ box-sizing: border-box; }}
      body {{ margin: 0; font-family: '{FONT}', sans-serif; color: #{INK}; font-size: 9.5pt; line-height: 1.35; }}
      .page {{ break-after: page; min-height: 255mm; position: relative; }}
      .page:last-child {{ break-after: auto; }}
      h1 {{ color: #{NAVY}; font-size: 20pt; margin: 0 0 7mm; border-bottom: 2px solid #{TEAL}; padding-bottom: 2mm; }}
      h2 {{ color: #{BLUE}; font-size: 14pt; margin: 0 0 4mm; }}
      h3 {{ color: #{TEAL}; font-size: 11pt; margin: 5mm 0 2mm; }}
      p {{ margin: 0 0 3mm; }}
      ul, ol {{ margin: 2mm 0 4mm 6mm; padding-left: 5mm; }}
      li {{ margin: 1mm 0; }}
      table {{ width: 100%; border-collapse: collapse; margin: 3mm 0 5mm; table-layout: fixed; }}
      th, td {{ border: 0.35mm solid #B7C0CA; padding: 2mm; vertical-align: top; word-wrap: break-word; }}
      thead th, table.summary tr:first-child th {{ background: #{NAVY}; color: white; font-weight: 700; }}
      tr:nth-child(even) td {{ background: #{PALE_GRAY}; }}
      .meta th {{ width: 20%; background: #{NAVY}; color: white; }}
      .meta td {{ width: 30%; background: #{PALE_GRAY}; }}
      .fields th {{ width: 25%; background: #{PALE_BLUE}; color: #{NAVY}; text-align: left; }}
      .fields td {{ background: white !important; }}
      .blank {{ min-height: 12mm; }}
      .blank.tall {{ min-height: 20mm; }}
      .callout {{ padding: 4mm; background: #{PALE_BLUE}; border-left: 2mm solid #{BLUE}; margin: 4mm 0; }}
      .warning {{ background: #{PALE_AMBER}; border-left-color: #{AMBER}; }}
      .safe {{ background: #{PALE_TEAL}; border-left-color: #{TEAL}; }}
      .cover {{ padding-top: 25mm; }}
      .brand {{ color: #{TEAL}; font-size: 17pt; font-weight: 700; letter-spacing: .5mm; }}
      .cover h1 {{ border: 0; font-size: 30pt; margin: 3mm 0; }}
      .subtitle {{ font-size: 13pt; color: #{MUTED}; margin-bottom: 12mm; }}
      .version {{ margin-top: 8mm; color: #{MUTED}; }}
      .secret {{ margin-top: 30mm; color: #{RED}; font-weight: 700; }}
      .summary {{ font-size: 8pt; }}
      .summary th:nth-child(1) {{ width: 9%; }} .summary th:nth-child(2) {{ width: 16%; }}
      .summary th:nth-child(3) {{ width: 24%; }} .summary th:nth-child(4) {{ width: 24%; }}
      .summary th:nth-child(5) {{ width: 10%; }} .summary th:nth-child(6) {{ width: 17%; }}
      .check {{ font-size: 16pt; text-align: center; }}
      .form .fields th {{ width: 25%; }}
      .small {{ font-size: 8pt; color: #{MUTED}; }}
    """
    intro = """
    <section class="page cover">
      <div class="brand">RAOS</div><h1>人手作業支援パック</h1>
      <div class="subtitle">決める・提供する・承認する作業を、Codexと分担して前へ進める</div>
      <div class="callout warning"><b>この資料の位置づけ</b><br>現行Canonicalに基づく記入用の非Canonical支援資料です。記入しただけではOpen Decisionは解決されません。人間の承認、Canonicalとの照合、所定のDecision RecordまたはEvidence登録が完了して初めて実装・本番判断へ使えます。</div>
      <div class="version">Version 1.0　|　2026-08-04　|　ja-JP</div>
      <div class="secret">重要: Secret、Password、Token、個人情報の原文は、この資料やチャットへ記載しないでください。</div>
    </section>
    <section class="page"><h1>1. 使い方</h1>
      <ol><li>人が事業・法務・契約・公開に関する選択を行います。</li><li>Codexが候補比較、リスク整理、設定案、チェックリスト、テストを作ります。</li><li>人がDecision Recordまたは承認票を確認し、承認します。</li><li>Codexが承認済み範囲だけを実装し、Local/CI/Runtime/Production証跡を分離して報告します。</li></ol>
      <div class="callout safe"><b>境界</b><br>Codexは、法的判断、規約同意、本人確認、契約締結、支払い承認、Secret値の保管、記事の最終公開、成果確定、Kill Switch解除を人間の代わりに行いません。</div>
      <h2>最初の30分で回答する5項目</h2>
      <table class="fields"><tr><th>1. 初期カテゴリ</th><td>決定済み: __________ / 候補3つ: ______________________________</td></tr><tr><th>2. ブランド・ドメイン</th><td>名称案: __________________ ドメイン候補: __________________</td></tr><tr><th>3. 予算・許容赤字</th><td>月額上限: ______円 許容赤字期間: ____か月 自動停止: ______円</td></tr><tr><th>4. Reviewer</th><td>主担当: ______ 代理: ______ 時間単価: ______円</td></tr><tr><th>5. 外部手続き</th><td>楽天: 未/済 Domain: 未/済 AWS: 未/済 Legal相談先: 未/済</td></tr></table>
      <h3>チャット回答テンプレート</h3><div class="callout">初期カテゴリ候補: / ブランド名候補: / ドメイン候補: / 月額予算上限: / 許容赤字期間: / 主Reviewer: / 代理Reviewer: / 標準時間単価: / 楽天Affiliate登録状況: / 法務相談先の有無:</div>
    </section>
    """
    summary = f"""
    <section class="page"><h1>2. Open Decision全体表</h1><p>15件のうち14件がProductionまたはGate blockerです。OD-004だけは手動Importで先へ進められます。</p>
      <table class="summary"><tr><th>ID</th><th>必要時点</th><th>判断・Evidence</th><th>Owner</th><th>Block</th><th>状態</th></tr>{summary_rows}</table>
      <h2>進行順序</h2><ul><li>A. 今決める: OD-001、002、005、008、009</li><li>B. カテゴリ設計時: OD-006、007</li><li>C. Finance/Live Adapter時: OD-003、015</li><li>D. Staging/Production前: OD-010、011、012、013、014</li><li>E. GATE-2高度化時: OD-004。未決でもSearch Consoleと手動CSVで継続可能</li></ul>
    </section>
    """
    forms = f"""
    <section class="page"><h1>4. 外部手続きチェックリスト</h1><p>契約・本人確認・支払い・媒体登録はAccount Owner本人が行います。Evidence参照には確認番号やファイル参照だけを記録し、Secret値を記載しません。</p>
      <table class="summary"><tr><th>完了</th><th>対象</th><th>人が行うこと</th><th>Owner</th><th>状態</th><th>Evidence参照</th></tr>{procedure_rows}</table>
      <div class="callout warning"><b>Credentialの扱い</b><br>Codexへ渡すのはSecret値ではなく、RAOS_RAKUTEN_ACCESS_KEYなどの参照名、権限範囲、有効期限、Rotation状態です。漏えいが疑われる場合、ローカル削除とProvider側失効は別作業として両方確認します。</div>
    </section>
    <section class="page form"><h1>5. 記事レビュー・公開承認票</h1><p>CodexはDraft、Evidence照合、Finding、Preview差分を準備します。ReviewerとApproverが内容を確認し、公開可否を決めます。</p><table class="fields">
      <tr><th>Article ID / Version</th><td></td></tr><tr><th>カテゴリ / Primary Decision</th><td></td></tr><tr><th>Reviewer / Approver</th><td>Reviewer: __________________　Approver: __________________</td></tr>
      <tr><th>Evidence</th><td>☐ 主要Claim 100%　☐ 全検証可能Claim 95%以上　☐ Source原本を確認</td></tr><tr><th>商品同定</th><td>☐ 型番　☐ 容量　☐ 色　☐ セット　☐ JAN　☐ 曖昧候補なし/処理済み</td></tr>
      <tr><th>品質</th><td>☐ 85点以上　☐ 各軸Floor合格　☐ 架空体験なし　☐ Tradeoff/対象外を明記</td></tr><tr><th>Policy</th><td>☐ Blocking Finding 0　☐ 広告表示　☐ 楽天遷移明示　☐ 画像/引用許可</td></tr>
      <tr><th>鮮度・リンク</th><td>☐ 価格　☐ 在庫　☐ 仕様　☐ 画像　☐ Affiliate linkが許容Age内</td></tr><tr><th>Preview差分</th><td>☐ 承認対象Version/Hashと一致　☐ 公開Renderer差分確認</td></tr>
      <tr><th>判断</th><td>☐ APPROVE　☐ REJECT　☐ PAUSE　理由: ______________________________</td></tr><tr><th>承認</th><td>Reviewer署名/ID: __________ 日付: ______　Approver署名/ID: __________ 日付: ______</td></tr></table>
      <div class="callout safe"><b>役割分離</b><br>編集者本人だけで最終承認・公開を完結させません。Critical AI TaskのReleaseは二人承認を要求します。</div>
    </section>
    <section class="page form"><h1>6. 月次Finance Review票</h1><p>発生報酬ではなく、楽天の確定成果と実費を用いて月次貢献利益を承認します。</p><table class="fields">
      <tr><th>対象月 / Report期間</th><td></td></tr><tr><th>Report Evidence</th><td>取得元・ファイルHash・取込ID:</td></tr><tr><th>照合</th><td>☐ 発生/確定を分離　☐ Cancel/否認/調整を確認　☐ 未帰属を分離　☐ 重複なし</td></tr>
      <tr><th>確定成果報酬</th><td>________________ 円</td></tr><tr><th>人件費</th><td>________________ 円　根拠:</td></tr><tr><th>AI/API/Cloud費</th><td>________________ 円</td></tr><tr><th>その他直接費</th><td>________________ 円</td></tr><tr><th>確定貢献利益</th><td>________________ 円</td></tr>
      <tr><th>未解決差異</th><td>金額: ______円　理由/調査Owner:</td></tr><tr><th>判断</th><td>☐ ACCEPT　☐ HOLD　☐ REIMPORT　☐ INVESTIGATE</td></tr><tr><th>承認</th><td>Finance Reviewer: __________________　日付: ____ / ____ / ____</td></tr></table>
    </section>
    <section class="page form"><h1>7. Release・Production承認票</h1><p>Local PASS、CI、Staging Runtime、Production readinessを混同せず、各OwnerがEvidenceを確認します。</p><table class="fields">
      <tr><th>Release ID / Commit</th><td></td></tr><tr><th>対象環境 / Scope</th><td></td></tr><tr><th>Test Evidence</th><td>☐ Contract　☐ Unit　☐ Integration　☐ E2E　☐ Security　☐ Manual</td></tr>
      <tr><th>Operations</th><td>☐ Rollback tested　☐ Backup restore current　☐ Alerts tested　☐ Runbooks current</td></tr><tr><th>Open Decisions</th><td>☐ Blocking Open Decision 0　未解決参照:</td></tr>
      <tr><th>Engineering</th><td>☐ APPROVE　☐ REJECT　Approver/Date:</td></tr><tr><th>Security</th><td>☐ APPROVE　☐ REJECT　Approver/Date:</td></tr><tr><th>Operations</th><td>☐ APPROVE　☐ REJECT　Approver/Date:</td></tr><tr><th>Product Owner</th><td>☐ APPROVE　☐ REJECT　Approver/Date:</td></tr><tr><th>最終判断</th><td>☐ NOT_READY　☐ CANARY　☐ PRODUCTION　☐ ROLLBACK</td></tr></table>
    </section>
    <section class="page form"><h1>8. Incident・Kill Switch承認票</h1><p>誤公開、Credential侵害、Affiliate不正、重大品質事故ではTraffic維持より安全停止を優先します。</p><table class="fields">
      <tr><th>Incident ID / Severity</th><td></td></tr><tr><th>検知日時 / 宣言者</th><td></td></tr><tr><th>影響</th><td>公開 / Affiliate / Data / Credential / Finance / その他:</td></tr><tr><th>初動</th><td>☐ 公開Kill　☐ Affiliate Kill　☐ Credential失効　☐ Rollback　☐ Provider連絡</td></tr>
      <tr><th>Evidence</th><td>Timeline/Log/Artifact参照:</td></tr><tr><th>復旧条件</th><td><div class="blank"></div></td></tr><tr><th>再開前確認</th><td>☐ Root cause　☐ Corrective action　☐ Regression test　☐ Owner sign-off</td></tr><tr><th>解除承認</th><td>Product Owner/Operator: __________________　日付: ____ / ____ / ____</td></tr></table>
    </section>
    <section class="page"><h1>9. 継続運用カレンダー</h1><p>Codexが収集・集計・差分作成を自動化し、人間が意味判断と承認を行います。</p><table><tr><th>周期</th><th>人が確認すること</th><th>主担当</th><th>実施/次回</th></tr>{operation_rows}</table>
      <h2>最小の人員構成</h2><ul><li>Product/Business Owner: 事業目標、カテゴリ、予算、契約、Gate判断</li><li>Managing Editor/Reviewer: 商品同定、Evidence、記事品質、推薦、最終編集</li><li>Operations/Security Ownerと代理: Account、Credential、Alert、Incident、Recovery</li><li>必要時のLegal/Privacy/Tax専門家: 規約、表示、Copyright、Privacy、税務</li></ul>
      <div class="callout safe"><b>最低人数</b><br>公開フローでは役割分離のため、少なくとも2つの独立した人間主体を確保します。本番運用では一次対応者と代理者も決めます。</div>
      <h2>Codexへの依頼テンプレート</h2><div class="callout">OD-___について支援してください。候補/前提は____です。比較軸は____、月額上限は____円、避けたいリスクは____です。Canonicalと公式情報を確認し、推奨案、代替案、未解決点、Decision Record案を作ってください。Secret値は扱わないでください。</div>
    </section>
    <section class="page"><h1>10. 参照元と注意事項</h1><p>この支援パックは次のRepository内Artifactを要約しています。Canonical/Upstreamは編集せず、差異がある場合はCanonicalの優先順位に従います。</p><ul>{source_items}</ul>
      <div class="callout warning"><b>完了の定義</b><br>チェック欄が埋まったこと自体はProduction readinessではありません。承認済みDecision、正式Evidence、実Runtime結果、Security/Operations/Product OwnerのSign-offがそろって初めてGate判定できます。</div><p class="small">End of document</p>
    </section>
    """
    document = f"<!doctype html><html lang='ja'><head><meta charset='utf-8'><title>RAOS 人手作業支援パック</title><style>{css}</style></head><body>{intro}{summary}{decision_pages}{forms}</body></html>"
    HTML_PATH.write_text(document, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_html()
    print(DOCX_PATH)
    print(HTML_PATH)
